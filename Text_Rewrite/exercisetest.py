import bs4
import requests
from bs4 import BeautifulSoup as bs
from urllib import request
import random
import re
import os
import urllib
from nltk.corpus import wordnet
from datamuse import datamuse
from docx import Document
from googletrans import Translator
import csv

n = 0
id_list = ['gexv2row1', 'gexv2row2']
doc_path = "C:/Users/Andrew/Documents/test.docx"
my_doc = Document(doc_path)

for i in range(4):
    with open('C:/Users/Andrew/Documents/Year4list.csv', 'r') as file:
        start = i * 20
        finish = start + 19
        print(start)
        print(finish)
        reader_rows = csv.reader(file.readlines()[start:finish])
        para = my_doc.add_paragraph('Words:  ')
        word_list = []
        for row in reader_rows:
            word = ''.join(row)
            word_list.append(word)
            print(word)
            para.add_run(word)
            para.add_run(' | ')
            my_doc.save(doc_path)
        my_doc.save(doc_path)
        print(word_list)
        random.shuffle(word_list)
        print(word_list)
        for item in word_list:
            url_word = ''.join(item)
            url = "https://www.dictionary.com/browse/" + url_word
            page_html = request.urlopen(url).read()
            page_soup = bs(page_html, "html.parser")
            random_id = random.choice(id_list)
            my_divs = page_soup.find_all("p", {"class": 'one-click-content css-fr4dvi-Sentence e15kc6du2'})
            for my_div in my_divs:
                n = n + 1
                p_str = my_div.text.strip()
                question = p_str.replace(url_word, '______________')
                if n < 2:
                    print(question)
                    my_doc.add_paragraph(question)
                else:
                    break
            my_doc.save(doc_path)
