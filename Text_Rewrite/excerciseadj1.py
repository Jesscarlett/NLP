import bs4
import requests
from bs4 import BeautifulSoup as bs
from urllib import request
import urllib.request
import random
import re
import os
import urllib
from nltk.corpus import wordnet
from datamuse import datamuse
from docx import Document
from googletrans import Translator
import csv

doc_path = "C:/Users/Andrew/Documents/Vocab/Adjective-Stage2-Vocab-excercise.docx"
my_doc = Document(doc_path)

# 5 words in a group
for i in range(14):
    with open('C:/Users/Andrew/Documents/Vocab/AdjYear234list.csv', 'r') as file:
        start = i * 10
        n = 0
        finish = start + 10
        print(start)
        print(finish)
        word_list = csv.reader(file.readlines()[start:finish])
        my_doc.add_paragraph('------ Exercise ------')
        word_group = []
        # print the questions
        for item in word_list:
            url_word = ''.join(item)
            word_group.append(url_word)
            print(word_group)
            n = n + 1
            # url = "https://dictionary.cambridge.org/dictionary/english/" + url_word

            url = "https://tangorin.com/sentences?search=" + url_word
            hdr = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:61.0)'}
            req = urllib.request.Request(url, headers=hdr)
            page_html = urllib.request.urlopen(req).read()
            page_soup = bs(page_html, "html.parser")
            # my_divs = page_soup.find_all("div", {"class": 'examp dexamp'})
            my_divs = page_soup.find_all("dd", {"class": 's-en'})
            m = 0
            if len(my_divs) > 15:
                for my_div in my_divs:
                    p_str = my_div.text.strip()
                    m = m + 1
                    if m < 2:
                        question = p_str.replace(url_word, '______________')
                        print(n, '. ', question)
                        para = my_doc.add_paragraph(str(n))
                        para.add_run('. ')
                        para.add_run(question)
                        continue
                    else:
                        break
                my_doc.save(doc_path)

            else:
                url = "https://dictionary.cambridge.org/dictionary/english/" + url_word
                hdr = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:61.0)'}
                req = urllib.request.Request(url, headers=hdr)
                page_html = urllib.request.urlopen(req).read()
                page_soup = bs(page_html, "html.parser")
                my_divs = page_soup.find_all("div", {"class": 'examp dexamp'})
                m = 0
                for my_div in my_divs:
                    p_str = my_div.text.strip()
                    m = m + 1
                    if m < 2:
                        question = p_str.replace(url_word, '______________')
                        print(n, '. ', question)
                        para = my_doc.add_paragraph(str(n))
                        para.add_run('. ')
                        para.add_run(question)
                        continue
                    else:
                        break
                my_doc.save(doc_path)

        # print word group
        random.shuffle(word_group)
        my_doc.add_paragraph(' ')
        my_doc.add_paragraph('-- Make Your Sentences --')
        for item in word_group:
            group_word = ''.join(item)
            print(group_word)
            my_doc.add_paragraph(group_word)
            my_doc.add_paragraph('__________________________________________________________________________________________________________________ ')
        my_doc.add_paragraph(' ')
        my_doc.save(doc_path)
