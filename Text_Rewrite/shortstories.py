import bs4
import requests
from bs4 import BeautifulSoup as bs
from urllib import request
import urllib.request
import random
import re
import os
import urllib
from nltk.corpus import wordnet
from datamuse import datamuse
from docx import Document
from googletrans import Translator
import csv


doc_path = "D:/ScarlettBooks/Children-Stories/ChildrenStories.docx"
my_doc = Document(doc_path)

with open('D:/ScarlettBooks/Children-Stories/storylist.csv', 'r') as file:
    reader = csv.reader(file)
    for row in reader:
        url = ''.join(row)
        # url = "https://www.shortkidstories.com/story/mollys-little-red-wagon/"
        hdr = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:61.0)'}
        req = urllib.request.Request(url, headers=hdr)
        page_html = urllib.request.urlopen(req).read()
        page_soup = bs(page_html, "html.parser")
        my_divs = page_soup.find_all("div", {"class": 'allStories'})
        my_titles = page_soup.find_all("h1", {"class": 'archiveTitle'})

        for my_title in my_titles:
            title_str = my_title.text.strip()
            print(title_str)
            my_doc.add_heading(title_str, 3)
        my_doc.save(doc_path)

        for my_div in my_divs:
            my_ps = my_div.find_all("p")
            # print(i)
            for my_p in my_ps:
                p_str = my_p.text.strip()
                print(p_str)
                my_doc.add_paragraph(p_str)
        my_doc.add_paragraph(' ')
        my_doc.add_paragraph('QUESTIONS')
        my_doc.add_paragraph(' ')
        my_doc.save(doc_path)

