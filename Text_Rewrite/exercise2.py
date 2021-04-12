import bs4
import requests
from bs4 import BeautifulSoup as bs
from urllib import request
import urllib.request
import random
import re
import os
import urllib
from nltk.corpus import wordnet
from datamuse import datamuse
from docx import Document
from googletrans import Translator
import csv

doc_path = "C:/Users/Andrew/Documents/Vocab/AUS-Grade1n2-Vocab-excercise2.docx"
my_doc = Document(doc_path)

# 5 words in a group
for i in range(38):
    with open('C:/Users/Andrew/Documents/Vocab/Year1n2list.csv', 'r') as file:
        start = i * 10
        n = 1
        finish = start + 10
        print(start)
        print(finish)
        word_list = csv.reader(file.readlines()[start:finish])
        my_doc.add_paragraph('------ Match Word with Meaning ------')
        word_group = []
        # print the questions
        for item in word_list:
            url_word = ''.join(item)
            try:
                syn = wordnet.synsets(url_word)[0]
                print(syn.definition())
                my_doc.add_paragraph('_________________ ' + syn.definition())
                my_doc.add_paragraph(' ')
                my_doc.save(doc_path)
            except:
                pass
            word_group.append(url_word)
        random.shuffle(word_group)
        print(word_group)
        for word in word_group:
            para = my_doc.add_paragraph(str(n))
            para.add_run('. ')
            para.add_run(word)
            n = n + 1
        my_doc.add_paragraph(' ')
        my_doc.add_paragraph('Pick 2 words to make a sentence, you can use a connective such as and, but, etc')
        my_doc.add_paragraph('_________________________________________________________________________________________________________________________________________________')
        my_doc.add_paragraph(' ')
        my_doc.save(doc_path)

