from textblob import TextBlob
import nltk
from textblob import Word
import sys
import bs4
import requests
from bs4 import BeautifulSoup as bs
from urllib import request
import urllib.request
import random
import re
import os
import urllib
from nltk.corpus import wordnet
from datamuse import datamuse
from docx import Document
from googletrans import Translator
import csv
import glob

# CC coordinating conjunction
# CD cardinal digit
# DT determiner
# EX existential there (like: “there is” … think of it like “there exists”)
# FW foreign word
# IN preposition/subordinating conjunction
# JJ adjective ‘big’
# JJR adjective, comparative ‘bigger’
# JJS adjective, superlative ‘biggest’
# LS list marker 1)
# MD modal could, will
# NN noun, singular ‘desk’
# NNS noun plural ‘desks’
# NNP proper noun, singular ‘Harrison’
# NNPS proper noun, plural ‘Americans’
# PDT predeterminer ‘all the kids’
# POS possessive ending parent‘s
# PRP personal pronoun I, he, she
# PRP$ possessive pronoun my, his, hers
# RB adverb very, silently,
# RBR adverb, comparative better
# RBS adverb, superlative best
# RP particle give up
# TO to go ‘to‘ the store.
# UH interjection errrrrrrrm
# VB verb, base form take
# VBD verb, past tense took
# VBG verb, gerund/present participle taking
# VBN verb, past participle taken
# VBP verb, sing. present, non-3d take
# VBZ verb, 3rd person sing. present takes
# WDT wh-determiner which
# WP wh-pronoun who, what
# WP$ possessive wh-pronoun whose
# WRB wh-abverb where, when

doc_path = "D:/ScarlettBooks/Children-Stories/The-Secret-Zoo/TheSecretZoo-Reading-Vocab-exercise.docx"
my_doc = Document(doc_path)
b = 0
for f in glob.glob('D:/ScarlettBooks/Children-Stories/The-Secret-Zoo/*.txt'):
    with open(f, 'r', encoding="utf-8") as openfile:
        print(f)
        b = b + 1
    # open txt file
    # with open(r'C:\Users\Andrew\Documents\Vocab\blank\data.txt', encoding='utf8') as f:
        my_doc.add_heading('EXERCISE {}'.format(b), 3)
        my_doc.add_paragraph(' ')
        string_org = openfile.read()
        string = string_org
        string1 = string_org
        # string = "THERE was once upon a time two farmers, and their names were Hudden and Dudden. They had poultry in their yards, sheep on the uplands, and scores of cattle in the meadow land alongside the river. But for all that they weren’t happy, for just between their two farms there lived a poor man by the name of Donald O’Neary. He had a hovel over his head and a strip of grass that was barely enough to keep his one cow, Daisy, from starving, and, though she did her best, it was but seldom that Donald got a drink of milk or a roll of butter from Daisy. You would think there was little here to make Hudden and Dudden jealous, but so it is, the more one has the more one wants, and Donald’s neighbors lay awake of nights scheming how they might get hold of his little strip of grass land."
        txt = TextBlob(string)
        tags = txt.tags
        adj_list = []
        adv_list = []
        verb_list = []
        t = 0
        v = 0
        n = 0
        w = 0
        # print(txt.word_counts)
        # get a list for adj, adv and verb
        for word, pos in tags:
            t = t + 1
            if pos == 'JJ' or pos == 'JJR' or pos == 'JJS':
                n = n + 1
                # print(word)
                if len(word) > 5:
                    adj_list.append(word)
            elif pos == 'VB' or pos == 'VBD' or pos == "VBG" or pos == "VBN" or pos == "VBP" or pos == "VBZ":
                v = v + 1
                # print(word)
                if len(word) > 5:
                    verb_list.append(word)
            elif pos == 'RB' or pos == 'NN' or pos == 'NNS':
                w = w + 1
                # print(word)
                if len(word) > 6:
                    adv_list.append(word)
        # print(t)
        # print(n)
        # print(v)
        # print(w)
        # print(verb_list)
        # print(adj_list)
        # print(adv_list)
        combined_list = verb_list + adv_list + adj_list
        # reduce duplicated words
        A_list = []
        for word in combined_list:
            count = combined_list.count(word)
            if count < 2 and (word.islower()) == True:
                A_list.append(word)
            else:
                print('duplication: ' + word)
        i = 0
        # for item in combined_list:
        #     if item not in final_list:
        #         final_list.append(item)
        # ensure only 10 words
        if len(A_list) > 10:
            random.shuffle(A_list)
            final_list = A_list[0:10]
        else:
            final_list = A_list
        print(final_list)
        # replace the word with lines and shuffle
        for item in final_list:
            if item in string:
                string = string.replace(item, '________________', 1)
        print(string)
        if '________________ ________________' in string:
            random.shuffle(A_list)
            final_list = A_list[0:10]
            print('2nd shuffle')
            for item1 in final_list:
                if item1 in string1:
                    string1 = string1.replace(item1, '________________', 1)
                # print(string)
            string = string1
        else:
            pass
        my_doc.add_paragraph(string)
        my_doc.add_paragraph(' ')
        random.shuffle(final_list)
        for text in final_list:
            my_doc.add_paragraph(text)
            my_doc.add_paragraph("_____________________________________________________________________")
        my_doc.add_paragraph(' ')
        my_doc.save(doc_path)
        # print(final_list)
        openfile.close()

