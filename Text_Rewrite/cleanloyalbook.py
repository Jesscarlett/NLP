import glob
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns
from docx2pdf import convert

# add page number


def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


# Folder Path
path = "D:/ScarlettBooks/booktest"
outpath = "D:/ScarlettBooks/outbooks"

# Change the directory
os.chdir(path)

# Read text File and replace with Powered by Jesscarlett
for file in os.listdir():
    # Check whether file is in text format or not
    file_path = f"{path}/{file}"
    try:
        with open(file_path, 'r', encoding="utf-8") as fin:
            infile = f"{file}.txt"
            outfile = f"{outpath}/{file}.txt"
            # first_line = fin.readline()
            # if first_line.isupper() == False:
            #     print(file)
            # else:
            #     continue
            delete_list = ["Provided by LoyalBooks.com"]
            with open(outfile, "w+", encoding="utf-8") as fout:
                for line in fin:
                    # if i < 1:
                    #     if line.isupper() == True:
                    #         fout.write(line)
                    #     else:
                    #         print(line)
                    #         continue
                    # else:
                    for word in delete_list:
                        line = line.replace(word, "POWERED BY JESSCARLETT")
                    fout.write(line)
        fout.close()
    except UnicodeEncodeError as error:
        print(error)
        print(infile)
        pass
# os.chdir(outpath)
# for f in glob.glob('D:/ScarlettBooks/outbooks/*.txt'):
#     new_name = f.replace(".txt", "")
#     os.rename(f, new_name)

# write Docx
for f in glob.glob('D:/ScarlettBooks/outbooks/*.txt'):
    try:
        with open(f, 'r', encoding="utf-8") as openfile:
            line = openfile.read()
            doc = Document()
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(12)
            doc.add_paragraph(line)
            # doc.save(f + ".docx")
            add_page_number(doc.sections[0].footer.paragraphs[0].add_run())
            doc.sections[0].footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.save(f + ".docx")
    except UnicodeDecodeError as error:
        print(error)
        print(f)
        pass

# remove .txt.txt in docx name
os.chdir(outpath)
for f in glob.glob('D:/ScarlettBooks/outbooks/*.docx'):
    new_name = f.replace(".txt.txt", "")
    os.rename(f, new_name)

# remove txt files
for t in glob.glob('D:/ScarlettBooks/outbooks/*.txt'):
    os.remove(t)

# convert to pdf
convert("D:/ScarlettBooks/outbooks/")

