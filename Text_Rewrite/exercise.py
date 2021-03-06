import bs4
import requests
from bs4 import BeautifulSoup as bs
from urllib import request
import urllib.request
import random
import re
import os
import urllib
from nltk.corpus import wordnet
from datamuse import datamuse
from docx import Document
from googletrans import Translator
import csv

doc_path = "C:/Users/Andrew/Documents/Vocab/AUS-Grade3-Vocab-excercise-answers.docx"
my_doc = Document(doc_path)

# 5 words in a group
for i in range(28):
    with open('C:/Users/Andrew/Documents/Vocab/Year3list.csv', 'r') as file:
        start = i * 5
        n = 0
        finish = start + 5
        print(start)
        print(finish)
        word_list = csv.reader(file.readlines()[start:finish])
        my_doc.add_paragraph('------ Exercise ------')
        word_group = []
        # print the questions
        for item in word_list:
            url_word = ''.join(item)
            word_group.append(url_word)
            print(word_group)
            n = n + 1
            # url = "https://www.dictionary.com/browse/" + url_word
            # page_html = request.urlopen(url).read()
            # page_soup = bs(page_html, "html.parser")
            # my_divs = page_soup.find_all("p", {"class": 'one-click-content css-fr4dvi-Sentence e15kc6du2'})

            url = "https://www.wordhippo.com/what-is/sentences-with-the-word/" + url_word + ".html"
            hdr = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:61.0)'}
            req = urllib.request.Request(url, headers=hdr)
            page_html = urllib.request.urlopen(url).read()
            page_soup = bs(page_html, "html.parser")
            my_divs = page_soup.find_all("tr", {"id": 'gexv2row2'})

            for my_div in my_divs:
                p_str = my_div.text.strip()
                question = p_str.replace(url_word, '______________')
                print(n, '. ', question)
                para = my_doc.add_paragraph(str(n))
                para.add_run('. ')
                para.add_run(question)
            my_doc.save(doc_path)
        # print word group
        random.shuffle(word_group)
        my_doc.add_paragraph(' ')
        my_doc.add_paragraph('-- Make Your Sentences --')
        for item in word_group:
            group_word = ''.join(item)
            print(group_word)
            my_doc.add_paragraph(group_word)
            try:
                syn = wordnet.synsets(group_word)[0]
                print('[Meaning] - ' + syn.definition())
                my_doc.add_paragraph('[Meaning] - ' + syn.definition())
                my_doc.save(doc_path)
            except:
                pass
            my_doc.add_paragraph('__________________________________________________________________________________________________________________ ')
        my_doc.add_paragraph(' ')
        my_doc.save(doc_path)
