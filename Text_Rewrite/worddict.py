import bs4
import requests
from bs4 import BeautifulSoup as bs
from urllib import request
import re
import os
import urllib
from nltk.corpus import wordnet
from datamuse import datamuse
from docx import Document
import csv
from PyDictionary import PyDictionary
dictionary = PyDictionary()
api = datamuse.Datamuse()

doc_path = "C:/Users/Andrew/Documents/Vocab/Adjective-Stage2-Vocab.docx"

with open('C:/Users/Andrew/Documents/Vocab/AdjYear234list.csv', 'r') as file:
    reader = csv.reader(file)
    for row in reader:
        word = ''.join(row)
        print(word)
        my_doc = Document(doc_path)
        try:
            # get word
            word_form = wordnet.synsets(word)[0].pos()
            print(word + ":[" + word_form + ".]")
            my_doc.add_heading(word, 3)

            # get part of speech
            dictionary = PyDictionary(word)
            val_list = dictionary.getMeanings().values()
            for val in val_list:
                print(val)
                print(type(val))
                parts = val.keys()
                print(parts)
                # para = my_doc.add_paragraph('[Part-of-Speech] - ')
                for part in parts:
                    print(part)
                    # para.add_run(part)
                    # para.add_run(' ')
                    break
            my_doc.save(doc_path)
        except:
            pass
        try:
            # get meaning
            syn = wordnet.synsets(word)[0]
            print('[Meaning] - ' + syn.definition())
            my_doc.add_paragraph('[Meaning] - ' + syn.definition())
            my_doc.save(doc_path)
        except:
            pass

        try:
            # get examples
            print('[Examples]')
            my_doc.add_paragraph('[Examples]')

            url = "https://sentence.yourdictionary.com/" + word
            hdr = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:61.0)'}
            req = urllib.request.Request(url, headers=hdr)
            page_html = request.urlopen(url).read()
            page_soup = bs(page_html, "html.parser")
            my_divs = page_soup.find_all("div", {"class": "sentence-item"})
            n = 0
            for div in my_divs:
                div_str = str(div.text)
                n = n + 1
                if n < 3:
                    print(div_str)
                    my_doc.add_paragraph(div_str)
                    my_doc.save(doc_path)
                    continue
                else:
                    break
        except:
            pass

        try:
            # get synonyms
            print('[Synonyms]')
            my_doc.add_paragraph('[Synonyms]')

            url2 = "https://thesaurus.yourdictionary.com/" + word
            hdr = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:61.0)'}
            req = urllib.request.Request(url, headers=hdr)
            page_html2 = request.urlopen(url2).read()
            page_soup2 = bs(page_html2, "html.parser")
            my_syns = page_soup2.find_all("a", {"class": "synonym-link"})
            x = 1
            syn_list = []
            for syn in my_syns:
                syn_str = str(syn.text)
                x = x + 1
                if x < 4:
                    syn_list.append(syn_str + "  ")
                    continue
                else:
                    print(syn_list)
                    my_doc.add_paragraph(syn_list)
                    my_doc.add_paragraph("[Copy and Learn]")
                    my_doc.add_paragraph("____________________________________________________________________________________________________________")
                    my_doc.add_paragraph(" ")
                    my_doc.save(doc_path)
                    break
        except:
            pass
