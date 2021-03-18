import bs4
import requests
from bs4 import BeautifulSoup as bs
from urllib import request
import random
import re
import os
import urllib
from nltk.corpus import wordnet
from datamuse import datamuse
from docx import Document
from googletrans import Translator
import csv

id_list = ['gexv2row1', 'gexv2row2']
doc_path = "C:/Users/Andrew/Documents/AUS-Grade4-Vocab-excercise.docx"
n = 0
with open('C:/Users/Andrew/Documents/Year4list.csv', 'r') as file:
    reader = csv.reader(file)
    for row in reader:
        word = ''.join(row)
        n = n + 1
        my_doc = Document(doc_path)
        try:
            url = "https://www.wordhippo.com/what-is/sentences-with-the-word/" + word + ".html"
            page_html = request.urlopen(url).read()
            page_soup = bs(page_html, "html.parser")
            random_id = random.choice(id_list)
            my_divs = page_soup.find_all("tr", {"id": random_id})
            for my_div in my_divs:
                td_str = my_div.text.strip()
                question = td_str.replace(word, '______________')
                print(question)
                my_doc.add_paragraph(question)
                my_doc.add_paragraph(' ')
                my_doc.save(doc_path)
                break

        except:
                pass