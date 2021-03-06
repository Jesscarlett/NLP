import bs4
import requests
from bs4 import BeautifulSoup as bs
from urllib import request
import urllib.request
import random
import re
import os
import urllib
from nltk.corpus import wordnet
from datamuse import datamuse
from docx import Document
from googletrans import Translator
import csv


class MathQuestion:

    def single_addition(self):
        r1 = random.randint(1, 50)
        r2 = random.randint(0, 49)
        list1 = random.choice(['+', '-'])
        return f'{r1 + r2} {list1} {r2} = __________'

    def double_addition(self):
        r1 = random.randint(100, 999)
        r2 = random.randint(10, 999)
        return f'{r1} + {r2} = __________'

    def addition1(self):
        r1 = random.randint(0, 79)
        r2 = random.randint(50, 99)
        r3 = random.randint(0, 60)
        return f'{r1} + {r2} = _____ + {r3}'

    def addition2(self):
        r1 = random.randint(0, 79)
        r2 = random.randint(50, 99)
        r3 = random.randint(0, 60)
        return f'{r1} + {r2} = {r3} + _____ '

    def addition3(self):
        r1 = random.randint(0, 79)
        r2 = random.randint(50, 99)
        r3 = random.randint(0, 60)
        return f'{r3} + _____ = {r1} + {r2}'

    def addition4(self):
        r1 = random.randint(0, 79)
        r2 = random.randint(50, 99)
        r3 = random.randint(0, 60)
        return f'_____ + {r3} = {r1} + {r2}'

    def subtraction1(self):
        r1 = random.randint(50, 100)
        r2 = random.randint(0, 30)
        r3 = random.randint(0, 30)
        return f'{r1} - _____  = {r2} + {r3}'

    def subtraction2(self):
        r1 = random.randint(50, 100)
        r2 = random.randint(0, 30)
        r3 = random.randint(0, 30)
        return f'{r2} + {r3} = {r1} - _____'

    def subtraction3(self):
        r1 = random.randint(0, 50)
        r2 = random.randint(0, 30)
        r3 = random.randint(0, 30)
        return f'{r2} + {r3} = _____ - {r1}'

    def multiplication(self):
        import random
        r1 = random.randint(1, 500)
        r2 = random.randint(1, 100)
        return f'{r1} x {r2} = ____'
        # print(f'{r1} x {r2} = ____')

    def devision(self):
        import random
        r1 = random.randint(20, 1000)
        r2 = random.randint(1, 12)
        return f'{r1} ?? {r2} = ____'

    def many_in(self):
        import random
        r1 = random.randint(20, 1000)
        r2 = random.randint(1, 12)
        return f'How many {r2} in {r1}? And how many left?'

    def shopping(self):
        import random
        list1 = ['cupcakes', 'pens', 'lolipops', 'partybags', 'bookmarks', 'rubbers', 'chalks', 'cheese-sticks',
                 'breakfast bars', 'markers', ' paperclips']
        list2 = [10, 15, 20, 25, 30, 50, 75, 100]
        list3 = ['Dad', 'Mummy', 'Jess', 'Scarlett', 'Lisa', 'Chloe', 'Emma', 'Bella', 'Benny', 'Lucy', 'Lola']
        list4 = ['the supermarket', 'Coles', 'the shop', 'Woolworth', 'Aldi', 'K-Mart', 'BigW', 'Myer', 'David Jones']
        item1 = random.choice(list1)
        r1 = random.randrange(5, 100, 5)
        r2 = random.randint(1, 50)
        item2 = random.choice(list1)
        r4 = random.randrange(5, 100, 5)
        r5 = random.randint(1, 50)
        r3 = random.choice(list2)
        name = random.choice(list3)
        shop = random.choice(list4)
        return f'{name} went to {shop}, she bought {r2} {item1}, each for {r1} cents, and then she bought {r5} {item2}, each for {r4} cents. She has {r3} dollars. How much change she got in the end?'

    def time_lap(self):
        import random
        list1 = ['Lisa', 'Chloe', 'Emma', 'Bella', 'Benny', 'Lucy', 'Lola', 'Jess', 'Scarlett']
        r1 = random.randint(6, 12)
        r2 = random.randint(9, 60)
        r3 = random.randint(1, 90)
        name = random.choice(list1)
        return f'{name} starts the test at {r1}:{r2}, it takes {r3} minutes. What time does {name} finish the test?'

    def time_trip(self):
        import random
        list1 = ['Lisa', 'Chloe', 'Emma', 'Bella', 'Benny', 'Lucy', 'Lola', 'Jess', 'Scarlett']
        r1 = random.randint(6, 12)
        r2 = random.randint(9, 60)
        r3 = random.randint(1, 90)
        name = random.choice(list1)
        return f'{name} leaves home for school at {r1}:{r2}, it takes {r3} minutes. What time does {name} get to school?'

    def fraction(self):
        import random
        r1 = random.randint(1, 6)
        r2 = random.randint(1, 10)
        r3 = random.randint(1, 12)
        r4 = r1 + r2
        r5 = (r1 + r2) * r3
        return f'{r1}/{r4} = ____ /{r5}'

    def time(self):
        import random
        list1 = ['half and hour', 'one quarter', 'one hour', '3 quarters', 'one and a half hour', '90 minutes',
                 'two hours', '3 and a quarter hours']
        list2 = ['before', 'after']
        r1 = random.randint(1, 13)
        r2 = random.randint(10, 59)
        r3 = random.choice(list1)
        r4 = random.choice(list2)
        return f'Now it is {r1}:{r2}, what is the time {r3} {r4} this?'

    def fraction_minus(self):
        import random
        list1 = ['a quarter', '3 quarters', '1/2', 'one and 1/2', 'one and 1/4', 'one and 3/4', '2 and 1/4', '2 and 1/2',
                 '2 and 3/4', '3 and a half']
        r1 = random.randint(4, 10)
        r2 = random.choice(list1)
        return f'What is {r1} take out {r2}?'

    def time_home(self):
        import random
        list1 = ['Lisa', 'Chloe', 'Emma', 'Bella', 'Benny', 'Lucy', 'Lola', 'Jess', 'Scarlett']
        r1 = random.randint(6, 12)
        r2 = random.randint(9, 60)
        r3 = random.randint(1, 90)
        name = random.choice(list1)
        return f'{name} arrives at Ballet class at {r1}:{r2}, she leaves home {r3} minutes before. What time does {name} leave home?'



    def oporder(self):
        import random
        list1 = ['+', '-']
        list2 = ['+', '-']
        list3 = ['x', '+']
        r1 = random.randint(30, 60)
        r2 = random.randint(11, 20)
        r3 = random.randint(1, 10)
        r4 = random.randint(1, 10)
        r5 = random.randint(5, 10)
        symbol1 = random.choice(list1)
        symbol2 = random.choice(list1)
        symbol3 = random.choice(list3)
        return f'{r1} + ({r2} {symbol1} {r3}) x {r4} {symbol2} ({r3} {symbol3} {r5}) = ______'

    def ampm(self):
        import random
        list1 = ['7:00', '7:15', '7:30', '7:45', '8:00', '8:15', '8:30', '8:45', '9:00', '9:10', '9:20']
        list2 = ['3:00', '3:15', '5:30', '4:45', '4:00', '5:15', '3:30', '5:45', '2:00', '2:10', '1:20']
        list3 = ['Scarlett', 'Jess', 'Chloe', 'Emma']
        r1 = random.choice(list1)
        r2 = random.choice(list2)
        r3 = random.choice(list3)
        return f'{r3} goes to school at {r1}AM, come back home at {r2}PM, how many hours do {r3} spend at school?'

i = ''
# p = Mmultiplication(self)
# a1 = MathQuestion.addition1(i)
# a2 = MathQuestion.addition2(i)
# a3 = MathQuestion.addition3(i)
# a4 = MathQuestion.addition4(i)
# a5 = MathQuestion.subtraction1(i)
# a6 = MathQuestion.subtraction2(i)
b = MathQuestion.time(i)
c = MathQuestion.ampm(i)
d = MathQuestion.oporder(i)
# e = MathQuestion.addition(i)
f = MathQuestion.time_home(i)
g = MathQuestion.time_lap(i)
h = MathQuestion.fraction_minus(i)
j = MathQuestion.time_trip(i)
k = MathQuestion.shopping(i)
l = MathQuestion.many_in(i)
m = MathQuestion.devision(i)


doc_path = "C:/Users/Andrew/Documents/Vocab/AUS-Stage1-Math.docx"
my_doc = Document(doc_path)

my_doc.add_heading('Subtraction II', 3)
my_doc.add_paragraph(' ')

for i in range(135):
    num = random.randint(1, 3)
    if num == 1:
        a5 = MathQuestion.subtraction1(i)
        print(a5)
        my_doc.add_paragraph(a5)
        my_doc.add_paragraph(' ')
    elif num == 2:
        a7 = MathQuestion.subtraction3(i)
        print(a7)
        my_doc.add_paragraph(a7)
        my_doc.add_paragraph(' ')
    else:
        a6 = MathQuestion.subtraction2(i)
        print(a6)
        my_doc.add_paragraph(a6)
        my_doc.add_paragraph(' ')

my_doc.add_page_break()
my_doc.save(doc_path)
# print(b)
# print(c)
# print(d)
# print(e)
# print(f)
# print(g)
# print(h)
# print(j)
# print(k)
# print(l)
# print(m)





