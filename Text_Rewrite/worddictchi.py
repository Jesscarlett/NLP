import bs4
import requests
from bs4 import BeautifulSoup as bs
from urllib import request
import re
import os
import urllib
from nltk.corpus import wordnet
from datamuse import datamuse
from docx import Document
from googletrans import Translator
import csv
api = datamuse.Datamuse()

doc_path = "C:/Users/Andrew/Documents/AUS-Grade4-Vocab-chinese.docx"

with open('C:/Users/Andrew/Documents/Year4list.csv', 'r') as file:
    reader = csv.reader(file)
    for row in reader:
        word = ''.join(row)
        my_doc = Document(doc_path)
        try:
            word_form = wordnet.synsets(word)[0].pos()

            print(word + ":[" + word_form + ".]")
            translator = Translator()
            translated = translator.translate(word, src='en', dest='zh-cn')
            my_doc.add_heading(word, 1)
            my_doc.add_paragraph(" ")

            syn = wordnet.synsets(word)[0]
            print('[Meaning] - ' + syn.definition())
            my_doc.add_paragraph('[Meaning] - ' + syn.definition() + '. ' + translated.text)

            print('[Examples]')
            my_doc.add_paragraph('[Examples]')

            url = "https://sentence.yourdictionary.com/" + word
            page_html = request.urlopen(url).read()
            page_soup = bs(page_html, "html.parser")
            my_divs = page_soup.find_all("div", {"class": "sentence-item"})
            n = 0
            for div in my_divs:
                div_str = str(div.text)
                n = n + 1
                if n < 4:
                    print(div_str)
                    my_doc.add_paragraph(div_str)
                    continue
                else:
                    break

            print('[Synonyms]')
            my_doc.add_paragraph('[Synonyms]')

            url2 = "https://thesaurus.yourdictionary.com/" + word
            page_html2 = request.urlopen(url2).read()
            page_soup2 = bs(page_html2, "html.parser")
            my_syns = page_soup2.find_all("a", {"class": "synonym-link"})
            x = 1
            syn_list = []
            for syn in my_syns:
                syn_str = str(syn.text)
                x = x + 1
                if x < 5:
                    syn_list.append(syn_str + "  ")
                    continue
                else:
                    print(syn_list)
                    my_doc.add_paragraph(syn_list)
                    my_doc.add_paragraph("-------------------------------------------------------------------")
                    my_doc.save(doc_path)
                    break
        except:
            pass
