import bs4
import requests
from bs4 import BeautifulSoup as bs
from urllib import request
import re
import os
import urllib
from nltk.corpus import wordnet
from datamuse import datamuse
from docx import Document
import csv
from PyDictionary import PyDictionary
from urllib.request import Request, urlopen
dictionary = PyDictionary()
api = datamuse.Datamuse()

doc_path = "D:/ScarlettBooks/Children-Stories/News-Reading-Comprehension5.docx"
# word = 'science'
# find urls
url = "https://www.kidsnews.com.au/red"
# page_html = request.urlopen(url).read()
req = Request(url, headers={'User-Agent': 'Mozilla/5.0'})
page_html = urlopen(req).read()
page_soup = bs(page_html, "html.parser")
my_divs = page_soup.find_all("h4", {"class": "heading"})
n = 0
urls = []
# print(my_divs)
for div in my_divs:
    cell = div.find('a')
    if cell == "" or cell is None:
        pass
    else:
        cell_link = cell.attrs.get("href")
        if cell_link == "" or cell_link is None:
            pass
        else:
            urls.append(cell_link)
print(urls)

# print articles
my_doc = Document(doc_path)
for page_url in urls[1:50]:
    req = Request(page_url, headers={'User-Agent': 'Mozilla/5.0'})
    page_html = urlopen(req).read()
    page_soup = bs(page_html, "html.parser")
    my_hs = page_soup.find_all("h1", {"class": "headline"})
    my_ps = page_soup.find_all("p", {"class": "capi-html"})
    n = 0
    # urls = []
    # print(my_hs)
    # print(my_ps)
    for h in my_hs:
        h_str = str(h.text)
        print(h_str)
        print(' ')
        my_doc.add_heading(h_str, 3)
        my_doc.add_paragraph(' ')
    for p in my_ps:
        p_str = str(p.text)
        print(' ')
        print(p_str)
        my_doc.add_paragraph(p_str)
        my_doc.save(doc_path)
        if p_str == 'GLOSSARY':
            break
        elif p_str == ' GLOSSARY':
            break
        else:
            continue
    my_doc.save(doc_path)

