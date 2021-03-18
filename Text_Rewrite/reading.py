import bs4
import requests
from bs4 import BeautifulSoup as bs
from urllib import request
import re
import os
import urllib
import csv
from docx import Document
from googletrans import Translator
import csv

doc_path = "D:/ScarlettBooks/temp/test.docx"
my_doc = Document(doc_path)
line_number = 0
n = 0
with open(r'D:/ScarlettBooks/temp/Anne of the Island.txt', 'r', encoding="utf8") as file:
    # contents = file.read()
    for line in file:
        line_number += 1
        if "Chapter I" in line:
            print(line_number)
            break
        else:
            continue
with open(r'D:/ScarlettBooks/temp/Anne of the Island.txt', 'r', encoding="utf8") as contents:
    for sentence in contents:
        n += 1
        if n < line_number:
            continue
        elif n > line_number + 70:
            if sentence == '\n':
                break
            else:
                print(sentence)
                string = sentence.rstrip("\n")
                my_doc.add_paragraph(string)
                my_doc.save(doc_path)
        else:
            if sentence != '\n':
                print(sentence)
                string = sentence.rstrip("\n")
                my_doc.add_paragraph(string)
                my_doc.save(doc_path)
            else:
                continue




